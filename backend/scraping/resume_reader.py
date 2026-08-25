"""
Fetch a candidate's resume file and pull the text out of it.

Fetch and extract only. Nothing here reads the rubric, scores anything, or
touches the evaluator -- the text is stored and, for now, nothing consumes it.

What a resume_link actually is
------------------------------
Not a file, usually. Surveyed across all 7,084 submissions holding a link on
2026-08-14:

    5,191  drive.google.com          73%
      721  docs.google.com           10%
      356  linkedin.com               5%    profile page, no file behind it
      106  1drv.ms                  1.5%
       69  dropbox.com              1.0%
       22  candidateassessments.ajaia.ai    the /apply page, not an upload

Two things follow, and both shape this module.

First, the portal hosts none of these. All 22 portal links are the application
page the candidate was looking at when they pasted. So the authenticated portal
session is irrelevant here and is not used: every fetch is an anonymous request
to a third-party host. The session is still threaded through fetch() because
the caller has one and a future portal-hosted upload would need it.

Second, none of these URLs is a file. They are share pages that answer with
HTML. Fetching them verbatim yields a Google sign-in page, not a PDF -- so
direct_url() rewrites the three hosts that matter into their download form.
Measured on live samples: Drive 24/30, Dropbox 11/12. Without the rewrite both
are zero.

Deliberately not chased: 1drv.ms (0/12 -- 403 on the share URL, 401 on the
OneDrive shares API), acrobat.adobe.com and canva.link (JS viewers), and
linkedin.com, which has no file to fetch at any URL. Together those are ~8% of
the corpus and each needs a browser, not a rewrite.

Expect roughly 40% of rows to end with a resume_error. That is the ceiling
without OCR or browser automation: ~20% of Drive files are private or deleted,
~5% of links are profile pages, ~4% point at a folder, and ~10% of the PDFs
that do arrive are scans with no text layer.
"""

import io
import logging
import re
from typing import Optional
from urllib.parse import parse_qsl, urlencode, urlparse, urlunparse

import requests

log = logging.getLogger(__name__)

# Enough of a resume to characterise a candidate. A CV that runs past this is
# padding or a portfolio dump, and the tail of it is the least informative part.
MAX_TEXT_CHARS = 8_000

# Below this, treat the extraction as having failed rather than succeeded.
#
# A scanned CV is not always cleanly empty. Observed in the smoke run: a
# photographed resume whose only text layer was the page numbers, extracting to
# "1 3 2 3" -- 11 characters. Stored as a success that would look like a
# candidate with a blank CV rather than a page nobody could read. The shortest
# genuine resume in the same sample ran to 1,900 characters, so 200 separates
# the two without argument.
MIN_TEXT_CHARS = 200

# A resume is a few hundred KB. Anything past this is a portfolio, a video, or
# a mis-pasted link, and is not worth pulling down to find that out.
MAX_BYTES = 20 * 1024 * 1024

FETCH_TIMEOUT = 30

# Sent on every request. Drive and Dropbox both serve a different, script-only
# page to a client that does not look like a browser.
USER_AGENT = "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"

# The file id in a Drive or Docs URL, in either of the two forms candidates
# paste: .../d/<id>/view and ...?id=<id>. 20 chars is below the shortest id
# observed (25) and above anything a path segment produces by accident.
_DRIVE_ID = re.compile(r"/d/([A-Za-z0-9_-]{20,})|[?&]id=([A-Za-z0-9_-]{20,})")


def _host(link: str) -> str:
    return (urlparse(link).netloc or "").lower().removeprefix("www.")


def direct_url(link: str) -> str:
    """
    Rewrite a share URL into the URL that serves the bytes.

    Unrecognised hosts are returned unchanged: a plain https://host/cv.pdf is
    already direct, and guessing at a viewer we have not measured would turn a
    clean 404 into a wrong-looking success.
    """
    host = _host(link)
    parsed = urlparse(link)

    if host in ("drive.google.com", "docs.google.com"):
        # A folder has no single file to download. Left alone so it fetches,
        # fails the type check, and is recorded as what it is.
        if "/folders/" in parsed.path:
            return link
        match = _DRIVE_ID.search(link)
        if not match:
            return link
        file_id = match.group(1) or match.group(2)
        if host == "docs.google.com" and "/document/" in parsed.path:
            # A native Google Doc has no stored file to download; it has to be
            # rendered to one. Uploaded .docx files live under this path too
            # and export the same way, so one branch covers both.
            return f"https://docs.google.com/document/d/{file_id}/export?format=pdf"
        return f"https://drive.google.com/uc?export=download&id={file_id}"

    if host == "dropbox.com":
        # dl=1 added to the existing query, never replacing it. The rlkey
        # parameter is the share token -- drop it and every link 404s.
        query = dict(parse_qsl(parsed.query))
        query["dl"] = "1"
        return urlunparse(parsed._replace(query=urlencode(query)))

    return link


def _sniff(data: bytes) -> Optional[str]:
    """The file type from its magic bytes, or None if it is not one we read."""
    if data.startswith(b"%PDF"):
        return "pdf"
    if data.startswith(b"PK\x03\x04"):
        # Every OOXML file is a zip. Only .docx has this part in it.
        return "docx" if b"word/" in data[:4096] or b"word/document.xml" in data else "zip"
    return None


def _type_from_content_type(content_type: str) -> Optional[str]:
    ctype = (content_type or "").split(";")[0].strip().lower()
    if ctype == "application/pdf":
        return "pdf"
    if ctype in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return "docx"
    if ctype.startswith("text/html"):
        return "html"
    return None


class FetchError(RuntimeError):
    """A resume that could not be retrieved, carrying the reason to store."""


# Failures that say nothing about the link and everything about the moment.
# Worth another attempt later; the rest are not.
#
# A private Drive file will be just as private tomorrow, and a LinkedIn profile
# will still have no file behind it -- retrying those is 1,400 pointless
# requests. A 429 is different: Google starts throttling around the 2,400th
# fetch of a backfill, and those rows are readable, just not right then.
TRANSIENT_ERRORS = (
    "http_429", "http_500", "http_502", "http_503", "http_504",
    "fetch_failed", "write_failed",
)


def is_transient(error: str) -> bool:
    """Whether a stored resume_error is worth another attempt."""
    return (error or "").startswith(TRANSIENT_ERRORS)


def fetch(link: str, session: Optional[requests.Session] = None) -> tuple[bytes, str]:
    """
    Download a resume. Returns (body, content_type), raises FetchError.

    `session` is accepted so the caller's authenticated portal session can be
    reused, but note that no resume observed is portal-hosted -- see the module
    docstring. It is only ever used for its connection pool today.
    """
    url = direct_url(link)
    getter = session or requests
    try:
        with getter.get(
            url,
            timeout=FETCH_TIMEOUT,
            stream=True,
            allow_redirects=True,
            headers={"User-Agent": USER_AGENT},
        ) as resp:
            if resp.status_code != 200:
                raise FetchError(f"http_{resp.status_code}")
            chunks, size = [], 0
            for chunk in resp.iter_content(chunk_size=1 << 16):
                chunks.append(chunk)
                size += len(chunk)
                if size > MAX_BYTES:
                    raise FetchError(f"too_large_over_{MAX_BYTES // 1024 // 1024}mb")
            return b"".join(chunks), resp.headers.get("Content-Type", "")
    except FetchError:
        raise
    except requests.RequestException as exc:
        raise FetchError(f"fetch_failed:{type(exc).__name__}") from exc


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    if reader.is_encrypted:
        # Empty-password decryption covers the common case: a PDF exported with
        # permissions set but no password to open it.
        try:
            if reader.decrypt("") == 0:
                raise FetchError("pdf_password_protected")
        except FetchError:
            raise
        except Exception as exc:
            raise FetchError("pdf_password_protected") from exc

    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # One malformed page should not cost the other five.
            continue
    return "\n".join(parts)


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Plenty of resumes are laid out as a borderless table, and those cells are
    # not paragraphs of the body -- skipping them loses the whole document.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _storable(text: str) -> str:
    """
    Drop characters that cannot survive the trip into MongoDB.

    A PDF's own encoding tables are frequently broken, and pypdf faithfully
    reports what it finds: half of a surrogate pair, most often the leading
    half of an emoji. Python holds a lone surrogate happily, but it is not
    valid UTF-8, so BSON refuses it and pymongo raises on the write.

    Found the hard way. The first full backfill died at row 2,264 of 3,692 on a
    single resume carrying a bare \\ud83d -- one candidate's CV ending the run
    for the 1,428 behind them. Cleaning here rather than at the call site keeps
    the guarantee where the text is produced: whatever this module returns can
    be written.

    NULs go too. Mongo tolerates them, but they are extraction debris and they
    truncate the text in most things that later display it.
    """
    cleaned = text.encode("utf-8", "replace").decode("utf-8")
    return cleaned.replace("\x00", "")


def _tidy(text: str) -> str:
    """Collapse the whitespace a PDF extractor leaves behind. Otherwise verbatim."""
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]
    out, blanks = [], 0
    for line in lines:
        if line.strip():
            blanks = 0
            out.append(line.strip())
        else:
            blanks += 1
            if blanks == 1 and out:
                out.append("")
    return "\n".join(out).strip()


def extract(data: bytes, content_type: str = "") -> str:
    """
    Text from a PDF or DOCX body. Raises FetchError if it is neither, or if it
    is one but holds no text.

    Type comes from the Content-Type header first and the magic bytes second,
    never the URL's extension -- candidates rename files, and the hosts that
    matter here serve every file under one generic type anyway (Dropbox sends
    application/binary for both PDF and DOCX).
    """
    kind = _type_from_content_type(content_type)
    if kind == "html":
        # A share page rather than a file: the link is private, expired, or was
        # never a file. Sniffed anyway, since some hosts mislabel a real PDF.
        kind = None
    if kind is None:
        kind = _sniff(data)

    if kind == "pdf":
        text = _extract_pdf(data)
    elif kind == "docx":
        text = _extract_docx(data)
    elif kind == "zip":
        raise FetchError("unsupported_type:zip")
    else:
        served = (content_type or "").split(";")[0].strip().lower() or "unknown"
        if data[:512].lstrip()[:9].lower().startswith((b"<html", b"<!doctype")):
            served = "text/html"
        raise FetchError(f"not_a_document:{served}")

    text = _tidy(_storable(text))
    if len(text) < MIN_TEXT_CHARS:
        # Overwhelmingly a scanned or image-only PDF: a photographed CV has no
        # text layer to extract, or has one holding nothing but page furniture.
        # Reading it needs OCR, which is a system binary and a different
        # decision. The length is reported so the threshold can be judged from
        # the stored errors rather than guessed at again.
        raise FetchError(f"no_text_extracted:{len(text)}_chars")
    return text[:MAX_TEXT_CHARS]


def read_resume(link: str,
                session: Optional[requests.Session] = None) -> tuple[str, str]:
    """
    Fetch and extract one resume. Returns (text, error) -- exactly one is set.

    Never raises. A resume that will not download is an ordinary outcome here,
    not a failure of the run: ~40% of links are private, deleted, a profile
    page, or a scan. The caller stores the reason and moves to the next one.
    """
    if not (link or "").strip():
        return "", "no_link"
    try:
        data, content_type = fetch(link, session)
        return extract(data, content_type), ""
    except FetchError as exc:
        return "", str(exc)
    except Exception as exc:                        # noqa: BLE001
        # A malformed PDF can raise almost anything out of a parser. One bad
        # file must not end a 3,700-row backfill.
        log.debug("Unexpected error reading %s: %s", link, exc)
        return "", f"extract_failed:{type(exc).__name__}"
